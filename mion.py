
from flask import Flask

app = Flask(__name__)

@app.route("/")
def home():
    return "Mion AI online!"
import os
import time
import threading
import datetime
import psutil
import pyttsx3
import speech_recognition as sr
import customtkinter as ctk
import webbrowser
import platform
import subprocess
import random
import string
import wave
import pyaudio
from tkinter import messagebox, filedialog, simpledialog

# Kutubxonalarni tekshirish
try:
    from docx import Document
    import pandas as pd
except ImportError:
    print("DIQQAT: 'python-docx' va 'pandas' kutubxonalarini o'rnating: pip install python-docx pandas openpyxl")

# =============================================================================
# MION NEURAL INTERFACE v6.0 - ULTIMATE NAVIGATOR EDITION
# =============================================================================

COLORS = {
    "bg": "#010103",
    "card": "#080812",
    "accent": "#00F2FF",      
    "highlight": "#BD00FF",   
    "office_blue": "#2B579A",
    "office_green": "#217346",
    "web_red": "#FF0000",
    "web_yellow": "#FFCC00",
    "danger": "#FF003C",      
    "success": "#00FF7F",
    "text": "#E0E0FF",
    "muted": "#404060"
}

class MionUltimate(ctk.CTk):
    def __init__(self):
        super().__init__()

        self.title("MION V6: ULTIMATE NAVIGATOR")
        self.geometry("1450x900")
        self.configure(fg_color=COLORS["bg"])
        
        # Audio & Tizim holati
        self.is_recording = False
        self.audio_frames = []
        self.engine = pyttsx3.init()
        self.setup_voice()
        self.recognizer = sr.Recognizer()

        self.setup_ui()
        self.start_monitoring()
        
        self.show_log("MION v6.0 Navigator yuklandi. Barcha modullar faol.", "SYSTEM")
        self.speak("Navigator tizimi tayyor. Qayerga boramiz?")

    def setup_voice(self):
        voices = self.engine.getProperty('voices')
        for voice in voices:
            if "female" in voice.name.lower():
                self.engine.setProperty('voice', voice.id)
                break
        self.engine.setProperty('rate', 180)

    def setup_ui(self):
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        # --- Sidebar (Scrollable) ---
        self.sidebar = ctk.CTkScrollableFrame(
            self, width=320, fg_color=COLORS["card"], 
            corner_radius=0, scrollbar_button_color=COLORS["muted"]
        )
        self.sidebar.grid(row=0, column=0, sticky="nsew")

        ctk.CTkLabel(self.sidebar, text="MION V6 PRO", font=("Orbitron", 26, "bold"), text_color=COLORS["accent"]).pack(pady=20)

        # SECTION: GLOBAL NAVIGATOR (YANGI)
        self.create_section_label("🌐 GLOBAL NAVIGATOR")
        self.add_side_btn("🔍 Google Qidiruv", lambda: self.open_web("https://google.com"), COLORS["text"])
        self.add_side_btn("📺 YouTube", lambda: self.open_web("https://youtube.com"), COLORS["web_red"])
        self.add_side_btn("🌍 Yandex", lambda: self.open_web("https://yandex.ru"), COLORS["web_yellow"])
        self.add_side_btn("🌐 Chrome Brauzer", lambda: self.launch_app("chrome"), COLORS["accent"])
        self.add_side_btn("🚀 Dasturni ishga tushirish", self.ask_and_launch, COLORS["highlight"])

        # SECTION: OFFICE MODULI
        self.create_section_label("📝 OFFICE & DOCUMENTS")
        self.add_side_btn("📘 Word-ga saqlash", self.export_to_word, COLORS["office_blue"])
        self.add_side_btn("📗 Excel-ga saqlash", self.export_to_excel, COLORS["office_green"])
        self.add_side_btn("📑 Wordni ochish", lambda: self.launch_app("winword"), COLORS["office_blue"])
        self.add_side_btn("📊 Excelni ochish", lambda: self.launch_app("excel"), COLORS["office_green"])

        # SECTION: DIKTOFON
        self.create_section_label("🎙️ AUDIO MODULI")
        self.rec_btn = self.add_side_btn("🔴 Yozishni boshlash", self.toggle_recording, COLORS["danger"])
        self.add_side_btn("📂 Audio arxiv", self.open_recordings, COLORS["text"])

        # SECTION: TIZIM MONITORI
        self.create_section_label("📊 TIZIM HOLATI")
        self.cpu_label, self.cpu_bar = self.create_monitor_item("CPU")
        self.ram_label, self.ram_bar = self.create_monitor_item("RAM")

        # EXIT
        self.add_side_btn("🛑 Tizimdan chiqish", self.quit_app, COLORS["danger"])

        # --- Main View ---
        self.main_view = ctk.CTkFrame(self, fg_color="transparent")
        self.main_view.grid(row=0, column=1, sticky="nsew", padx=25, pady=25)
        self.main_view.grid_rowconfigure(1, weight=1)
        self.main_view.grid_columnconfigure(0, weight=1)

        # Dashboard Header
        self.header_frame = ctk.CTkFrame(self.main_view, fg_color="transparent")
        self.header_frame.grid(row=0, column=0, sticky="ew", pady=(0, 20))
        
        self.status_indicator = ctk.CTkLabel(self.header_frame, text="● TIZIM ON-LINE", font=("Consolas", 12), text_color=COLORS["success"])
        self.status_indicator.pack(side="left")

        # Terminal
        self.terminal = ctk.CTkTextbox(
            self.main_view, font=("Consolas", 15), fg_color="#05050C", 
            text_color=COLORS["text"], border_width=1, border_color="#1A1A30"
        )
        self.terminal.grid(row=1, column=0, sticky="nsew", pady=(0, 20))

        # Input Area
        self.input_frame = ctk.CTkFrame(self.main_view, fg_color="transparent")
        self.input_frame.grid(row=2, column=0, sticky="ew")
        self.input_frame.grid_columnconfigure(0, weight=1)

        self.user_entry = ctk.CTkEntry(
            self.input_frame, placeholder_text="Buyruq bering (masalan: 'youtube och' yoki 'wordga yoz')...",
            height=60, font=("Segoe UI", 16), fg_color=COLORS["card"], border_color="#1A1A30"
        )
        self.user_entry.grid(row=0, column=0, sticky="ew", padx=(0, 10))
        self.user_entry.bind("<Return>", lambda e: self.process_text_input())

        self.voice_btn = ctk.CTkButton(
            self.input_frame, text="🎤", width=60, height=60, font=("Arial", 24),
            fg_color=COLORS["highlight"], command=self.toggle_voice_command
        )
        self.voice_btn.grid(row=0, column=1)

    # --- NAVIGATSIYA VA APP LAUNCHER ---

    def open_web(self, url):
        self.show_log(f"Veb-sahifa ochilmoqda: {url}", "NAVIGATOR")
        webbrowser.open(url)

    def ask_and_launch(self):
        # Dastur tanlash dialogi
        app_name = simpledialog.askstring("Launcher", "Ochmoqchi bo'lgan ilova nomini kiriting:\n(masalan: notepad, calc, chrome, telegram)")
        if app_name:
            self.launch_app(app_name)

    def launch_app(self, app_name):
        try:
            self.show_log(f"Ilova ishga tushirilmoqda: {app_name}", "LAUNCHER")
            # Windows uchun 'start' buyrug'i
            subprocess.Popen(f"start {app_name}", shell=True)
            self.speak(f"{app_name} ochilmoqda")
        except Exception as e:
            self.show_log(f"Xato: {app_name} ilovasini ochib bo'lmadi.", "ERROR")

    # --- UI YORDAMCHILARI ---

    def create_section_label(self, text):
        lbl = ctk.CTkLabel(self.sidebar, text=text, font=("Orbitron", 12, "bold"), text_color=COLORS["muted"])
        lbl.pack(pady=(20, 8), padx=15, anchor="w")

    def create_monitor_item(self, name):
        lbl = ctk.CTkLabel(self.sidebar, text=f"{name}: 0%", font=("Consolas", 12), text_color=COLORS["text"])
        lbl.pack(padx=20, anchor="w")
        bar = ctk.CTkProgressBar(self.sidebar, height=4, progress_color=COLORS["accent"], fg_color="#151525")
        bar.set(0)
        bar.pack(fill="x", padx=20, pady=(2, 10))
        return lbl, bar

    def add_side_btn(self, text, cmd, color):
        btn = ctk.CTkButton(
            self.sidebar, text=text, command=cmd, height=45,
            fg_color="transparent", border_width=1, border_color="#1A1A30",
            hover_color="#101020", text_color=color, anchor="w", font=("Consolas", 13)
        )
        btn.pack(fill="x", padx=15, pady=4)
        return btn

    # --- OFFICE & AUDIO (AVVALGI FUNKSIYALAR) ---

    def export_to_word(self):
        content = self.terminal.get("1.0", "end").strip()
        if not content: return
        doc = Document()
        doc.add_heading('MION Hisoboti', 0)
        doc.add_paragraph(content)
        fpath = filedialog.asksaveasfilename(defaultextension=".docx")
        if fpath: doc.save(fpath); self.show_log("Word saqlandi", "OFFICE")

    def export_to_excel(self):
        content = self.terminal.get("1.0", "end").strip().split('\n')
        df = pd.DataFrame({"Ma'lumot": content})
        fpath = filedialog.asksaveasfilename(defaultextension=".xlsx")
        if fpath: df.to_excel(fpath, index=False); self.show_log("Excel saqlandi", "OFFICE")

    def toggle_recording(self):
        if not self.is_recording:
            self.is_recording = True
            self.rec_btn.configure(text="⏹️ STOP", fg_color=COLORS["danger"])
            threading.Thread(target=self.record_audio_logic, daemon=True).start()
        else:
            self.is_recording = False
            self.rec_btn.configure(text="🔴 REC", fg_color="transparent")

    def record_audio_logic(self):
        p = pyaudio.PyAudio()
        stream = p.open(format=pyaudio.paInt16, channels=1, rate=44100, input=True, frames_per_buffer=1024)
        self.audio_frames = []
        while self.is_recording:
            self.audio_frames.append(stream.read(1024))
        stream.stop_stream(); stream.close(); p.terminate()
        if not os.path.exists("recordings"): os.makedirs("recordings")
        fname = f"recordings/REC_{datetime.datetime.now().strftime('%H%M%S')}.wav"
        wf = wave.open(fname, 'wb')
        wf.setnchannels(1); wf.setsampwidth(p.get_sample_size(pyaudio.paInt16)); wf.setframerate(44100)
        wf.writeframes(b''.join(self.audio_frames)); wf.close()
        self.show_log(f"Audio saqlandi: {fname}", "AUDIO")

    def open_recordings(self):
        if os.path.exists("recordings"): os.startfile("recordings")

    # --- SYSTEM CORE ---

    def show_log(self, text, sender="MION"):
        now = datetime.datetime.now().strftime("%H:%M:%S")
        self.terminal.insert("end", f"[{now}] [{sender}] > {text}\n\n")
        self.terminal.see("end")

    def speak(self, text):
        def _say():
            try: self.engine.say(text); self.engine.runAndWait()
            except: pass
        threading.Thread(target=_say, daemon=True).start()

    def start_monitoring(self):
        def _loop():
            while True:
                c, r = psutil.cpu_percent(), psutil.virtual_memory().percent
                self.after(0, lambda: self.update_monitors(c, r))
                time.sleep(2)
        threading.Thread(target=_loop, daemon=True).start()

    def update_monitors(self, c, r):
        self.cpu_bar.set(c/100); self.cpu_label.configure(text=f"CPU: {c}%")
        self.ram_bar.set(r/100); self.ram_label.configure(text=f"RAM: {r}%")

    def process_text_input(self):
        query = self.user_entry.get()
        if query:
            self.user_entry.delete(0, "end")
            self.process_query(query)

    def process_query(self, query):
        self.show_log(query, "USER")
        q = query.lower()
        if "google" in q: self.open_web("https://google.com")
        elif "youtube" in q: self.open_web("https://youtube.com")
        elif "chrome" in q: self.launch_app("chrome")
        elif "yandex" in q: self.open_web("https://yandex.ru")
        elif "och" in q or "open" in q:
            app = q.replace("och", "").replace("open", "").strip()
            self.launch_app(app)
        else: self.speak("Buyruq qabul qilindi.")

    def toggle_voice_command(self):
        self.speak("Sizni eshitaman")
        threading.Thread(target=self.listen_voice, daemon=True).start()

    def listen_voice(self):
        with sr.Microphone() as source:
            try:
                audio = self.recognizer.listen(source, timeout=3, phrase_time_limit=3)
                txt = self.recognizer.recognize_google(audio, language="uz-UZ")
                self.after(0, lambda: self.process_query(txt))
            except: pass

    def quit_app(self):
        self.destroy()

if __name__ == "__main__":
    ctk.set_appearance_mode("Dark")
    app = MionUltimate()
    app.mainloop()
